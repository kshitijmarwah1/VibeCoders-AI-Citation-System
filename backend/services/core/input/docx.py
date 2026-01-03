from docx import Document
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from DOCX file with error handling.
    """
    if not file_path or not Path(file_path).exists():
        raise FileNotFoundError(f"DOCX file not found: {file_path}")

    try:
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        
        if not paragraphs:
            raise ValueError(f"No text content found in DOCX: {file_path}")
            
        return "\n".join(paragraphs)
    except Exception as e:
        logger.error(f"Failed to extract text from DOCX {file_path}: {e}")
        raise ValueError(f"Invalid or corrupted DOCX file: {file_path}")
